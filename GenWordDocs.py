# Generate word documents
import os
import random
import string
import sys
from docx import Document


# Define the path where the Word documents will be saved
output_dir = "C:\\Test\\Documents"
os.makedirs(output_dir, exist_ok=True)

# Load a list of English words
with open("words.txt") as word_file:
    english_words = word_file.read().splitlines()

# Function to generate a random list of words
def generate_random_words(num_words):
    return random.sample(english_words, num_words)

# Function to generate metadata for a Word document
def generate_metadata(document, title, subject, author, keywords, comments):
    core_properties = document.core_properties
    core_properties.title = title
    core_properties.subject = subject
    core_properties.author = author
    core_properties.keywords = keywords
    core_properties.comments = comments
    
    
# Create 5000 Word documents
for i in range(1, 5001):
    document = Document()
    document.add_heading(f"Document {i}", level=1)
    
    # Add 5 paragraphs of 10 words each
    for _ in range(5):
        random_words = generate_random_words(10)
        random_text = ' '.join(random_words)
        document.add_paragraph(random_text)

    # Generate metadata for the document
    generate_metadata(
        document,
        title=f"Document {i}",
        subject="Generated Document",
        author="Automated Script",
        keywords="random, words, generated, document",
        comments="This document was generated automatically as part of a batch of x documents."
    )
      
    # Save the document
    file_path = os.path.join(output_dir, f"document_{i}.docx")
    document.save(file_path)

print("5000 Word documents have been created successfully.")
